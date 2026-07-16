from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from paper_agent.models import AuditReport, Draft, EvidenceCard


def export_markdown(
    draft: Draft,
    evidence: list[EvidenceCard],
    audit: AuditReport,
    output: Path,
) -> Path:
    references = ["## 来源台账"]
    evidence_map = {item.id: item for item in evidence}
    cited = []
    for evidence_id in re.findall(r"\[E:(E-[A-Za-z0-9-]+)\]", draft.markdown):
        if evidence_id in evidence_map and evidence_id not in cited:
            cited.append(evidence_id)
    for item_id in cited:
        item = evidence_map[item_id]
        references.append(
            f"- `{item.id}`：文档 `{item.document_id}`，位置 `{item.location}`。{item.summary}"
        )
    audit_lines = ["## 生成审计摘要", f"- 是否通过阻断检查：{'是' if audit.passed else '否'}"]
    for key, value in audit.metrics.items():
        audit_lines.append(f"- {key}: {value}")
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(
        draft.markdown + "\n" + "\n".join(references) + "\n\n" + "\n".join(audit_lines) + "\n",
        encoding="utf-8",
    )
    return output


def export_docx(draft: Draft, audit: AuditReport, output: Path) -> Path:
    document = DocxDocument()
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    title = document.add_heading(draft.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for section in draft.sections:
        document.add_heading(section.title, level=1)
        for raw in re.split(r"\n\s*\n", section.content.strip()):
            value = re.sub(r"^#{1,6}\s*", "", raw.strip())
            if value:
                paragraph = document.add_paragraph(value)
                paragraph.paragraph_format.first_line_indent = Pt(24)
                paragraph.paragraph_format.line_spacing = 1.5

    document.add_page_break()  # type: ignore[no-untyped-call]
    document.add_heading("生成审计摘要", level=1)
    document.add_paragraph(f"阻断检查：{'通过' if audit.passed else '未通过'}")
    for key, value in audit.metrics.items():
        document.add_paragraph(f"{key}: {value}")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output

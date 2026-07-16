from __future__ import annotations

import hashlib
import mimetypes
import re
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from paper_agent.models import Document, DocumentRole, TextBlock


class UnsupportedDocumentError(ValueError):
    pass


MAX_FILE_BYTES = 25 * 1024 * 1024
MAX_PDF_PAGES = 500
INJECTION_PATTERNS = (
    r"ignore\s+(?:all\s+)?previous\s+instructions",
    r"忽略(?:以上|此前|之前|所有)指令",
    r"system\s+prompt",
    r"你现在是.{0,20}(?:助手|agent|模型)",
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _blocks_from_plain(text: str, *, markdown: bool) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    current_heading: str | None = None
    index = 0
    for raw in re.split(r"\n\s*\n", text.replace("\r\n", "\n")):
        value = raw.strip()
        if not value:
            continue
        level: int | None = None
        kind = "paragraph"
        if markdown and (match := re.match(r"^(#{1,6})\s+(.+)$", value)):
            level = len(match.group(1))
            value = match.group(2).strip()
            current_heading = value
            kind = "heading"
        blocks.append(
            TextBlock(
                index=index,
                text=value,
                heading=current_heading,
                level=level,
                kind=kind,
            )
        )
        index += 1
    return blocks


def _parse_docx(path: Path) -> tuple[list[TextBlock], list[str]]:
    document = DocxDocument(path)
    blocks: list[TextBlock] = []
    current_heading: str | None = None
    warnings: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        level: int | None = None
        kind = "paragraph"
        if style.lower().startswith("heading") or style.startswith("标题"):
            digits = re.findall(r"\d+", style)
            level = int(digits[0]) if digits else 1
            current_heading = text
            kind = "heading"
        blocks.append(
            TextBlock(
                index=len(blocks),
                text=text,
                heading=current_heading,
                level=level,
                kind=kind,
            )
        )
    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(
                    TextBlock(
                        index=len(blocks),
                        text=" | ".join(cells),
                        heading=current_heading,
                        kind=f"table-{table_number}-row-{row_number}",
                    )
                )
    if not blocks:
        warnings.append("DOCX 未提取到正文；文件可能只包含图片或不支持的对象。")
    return blocks, warnings


def _parse_pdf(path: Path) -> tuple[list[TextBlock], list[str]]:
    reader = PdfReader(str(path))
    if len(reader.pages) > MAX_PDF_PAGES:
        raise UnsupportedDocumentError(f"PDF 超过 {MAX_PDF_PAGES} 页的首版安全限制。")
    blocks: list[TextBlock] = []
    warnings: list[str] = []
    empty_pages = 0
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            empty_pages += 1
            continue
        for part in re.split(r"\n\s*\n|(?<=。)\s*\n", text):
            value = re.sub(r"\s+", " ", part).strip()
            if value:
                blocks.append(
                    TextBlock(index=len(blocks), text=value, page=page_number, kind="paragraph")
                )
    if not blocks:
        raise UnsupportedDocumentError("PDF 未检测到文字层，可能是扫描版；首版暂不支持 OCR。")
    if empty_pages:
        warnings.append(f"{empty_pages} 页未提取到文字，请检查是否含扫描页或复杂排版。")
    return blocks, warnings


def ingest_document(path: Path, role: DocumentRole) -> Document:
    path = path.expanduser().resolve()
    if not path.is_file():
        raise FileNotFoundError(path)
    if path.stat().st_size > MAX_FILE_BYTES:
        raise UnsupportedDocumentError("文件超过 25 MB 的首版安全限制。")
    suffix = path.suffix.lower()
    warnings: list[str] = []
    if suffix in {".txt", ".md", ".markdown"}:
        text = path.read_text(encoding="utf-8-sig")
        blocks = _blocks_from_plain(text, markdown=suffix != ".txt")
    elif suffix == ".docx":
        blocks, warnings = _parse_docx(path)
    elif suffix == ".pdf":
        blocks, warnings = _parse_pdf(path)
    else:
        raise UnsupportedDocumentError(
            f"不支持 {suffix or '无扩展名'}；首版仅支持 DOCX、文字版 PDF、Markdown、TXT。"
        )
    if not blocks:
        raise UnsupportedDocumentError(f"文档没有可用文本：{path.name}")
    joined = "\n".join(block.text for block in blocks)
    if any(re.search(pattern, joined, re.IGNORECASE) for pattern in INJECTION_PATTERNS):
        warnings.append("检测到疑似 Prompt Injection；内容将只作为不可信资料处理。")
    title = next((b.text for b in blocks if b.kind == "heading"), path.stem)
    return Document(
        id=uuid.uuid4().hex[:12],
        role=role,
        source_path=str(path),
        filename=path.name,
        sha256=_sha256(path),
        media_type=mimetypes.guess_type(path.name)[0] or "application/octet-stream",
        title=title,
        blocks=blocks,
        warnings=warnings,
    )
